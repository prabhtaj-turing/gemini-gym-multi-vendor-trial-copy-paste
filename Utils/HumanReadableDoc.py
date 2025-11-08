#%%
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Optional, Any
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import json


# --- Constants ---
CODE_FONT_NAME: str = 'Courier New'
CODE_FONT_SIZE: Pt = Pt(10)
CODE_BLOCK_BACKGROUND_COLOR: str = 'F0F0F0'  # Light gray
DEFAULT_HEADING_LEVEL: int = 1 # Default docx level for generic headings
ASSUMED_SECTION_HEADING_LEVEL: int = 3  # H3 for "Args:", "Returns:"
MARKDOWN_HEADING_MAX_LEVEL: int = 4 # Cap #-based headings at docx H4 (as per original implicit logic)

LIST_INDENT_PER_LEVEL: Inches = Inches(0.25)
UNORDERED_LIST_MARKERS: List[str] = ["- ", "* ", "• "]
FENCED_CODE_BLOCK_MARKER: str = "```"
HEADING_MARKER: str = "#"

# Precompiled Regex Patterns
ORDERED_LIST_PATTERN: re.Pattern[str] = re.compile(r"^\d+\.\s+")
PYTHON_DEF_PATTERN: re.Pattern[str] = re.compile(r"^def\s+.*\):$") # Matches 'def function(...):'
SECTION_TITLE_KEYWORDS: tuple[str, str] = ("args:", "returns:", "raises:")


class MarkdownToDocxConverter:
    """
    Converts Markdown text content to a DOCX document.
    It handles common Markdown elements like headings, lists, code blocks,
    and paragraphs, with some heuristic-based conversions.
    """

    def __init__(self):
        """Initializes the converter with a new DOCX document instance."""
        self.doc: Document = Document()

    def _preprocess_line_content(self, text: str) -> str:
        """
        Preprocesses text content to handle Markdown escapes and strip specific HTML tags.
        - Removes backslashes used to escape Markdown characters.
        - Removes empty anchor tags like <a id="..."></a>.
        Args:
            text: The raw text string.
        Returns:
            The processed and stripped text string.
        """
        processed_text = text
        # Remove backslash escapes for Markdown special characters
        processed_text = re.sub(r'\\([\\`*_{}\[\]()#+-.!])', r'\1', processed_text)
        # Remove empty anchor tags (often from Markdown IDs)
        processed_text = re.sub(r'<a\s+id="[^"]*"\s*>\s*</a>', '', processed_text)
        return processed_text.strip()

    def _format_code_paragraph(self, paragraph: 'docx.paragraph.Paragraph', text_content: str, is_block: bool = True):
        """
        Applies styling to a paragraph intended for code.
        Sets font to Courier New, size 10pt.
        If is_block is True, also applies a light gray background shading.

        Args:
            paragraph: The docx.paragraph.Paragraph object to style.
            text_content: The string content to add as a run in the paragraph.
            is_block: If True, applies block-level styling (background).
        """
        run = paragraph.add_run(text_content)
        run.font.name = CODE_FONT_NAME
        run.font.size = CODE_FONT_SIZE

        if is_block:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            try:
                pPr = paragraph._p.get_or_add_pPr() # type: ignore[attr-defined] # Accessing protected member for styling
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), CODE_BLOCK_BACKGROUND_COLOR)
                pPr.append(shd)
            except Exception as e:
                print(f"Warning: Could not apply background shading for code block. Error: {e}")

    def _add_heading(self, text_content: str, markdown_level: Optional[int] = None, explicit_docx_level: Optional[int] = None):
        """
        Adds a heading to the document with appropriate styling.

        Args:
            text_content: The text of the heading (already stripped of # markers).
            markdown_level: The number of '#' characters (e.g., 1 for H1, 2 for H2).
            explicit_docx_level: A specific DOCX heading level to apply (1-9).
        """
        processed_text = self._preprocess_line_content(text_content)
        if not processed_text:
            return

        docx_level_to_apply: int = DEFAULT_HEADING_LEVEL

        if explicit_docx_level is not None:
            docx_level_to_apply = explicit_docx_level
        elif markdown_level is not None and markdown_level > 0:
            docx_level_to_apply = markdown_level # Map Markdown H1 to docx Level 1, etc.

        # Apply capping for #-based Markdown headings (max H4, as per original implicit behavior)
        if markdown_level is not None:
            docx_level_to_apply = min(docx_level_to_apply, MARKDOWN_HEADING_MAX_LEVEL)

        # Ensure the level is within docx's typical supported range (1-9 for headings)
        final_docx_level = max(1, min(docx_level_to_apply, 9))

        self.doc.add_heading(processed_text, level=final_docx_level)

    def _consume_fenced_code_block(self, all_lines: List[str], current_index: int) -> int:
        """
        Consumes lines belonging to a fenced code block (```) and adds them to the document.
        Args:
            all_lines: The list of all raw lines from the Markdown file.
            current_index: The index of the line where '```' was detected.
        Returns:
            The index of the line immediately after the processed code block.
        """
        code_lines_content: List[str] = []
        idx = current_index + 1  # Start reading from the line after the opening ```

        while idx < len(all_lines):
            line_content = all_lines[idx].rstrip('\n')
            if line_content.strip() == FENCED_CODE_BLOCK_MARKER:
                idx += 1  # Consume the closing ```
                break
            code_lines_content.append(line_content) # Preserve original line content for code
            idx += 1
        else:
            # Closing ``` not found before EOF
            print(f"Warning: Closing '{FENCED_CODE_BLOCK_MARKER}' not found for code block starting near line {current_index + 1}.")

        if code_lines_content:
            paragraph = self.doc.add_paragraph()
            self._format_code_paragraph(paragraph, "\n".join(code_lines_content), is_block=True)
        elif idx == current_index + 1 : # Empty code block ``` ```
            paragraph = self.doc.add_paragraph() # Add an empty but styled block
            self._format_code_paragraph(paragraph, "", is_block=True)

        return idx

    def _add_list_item(self, text: str, indent_level: int, is_ordered: bool = False, item_number: Optional[int] = None):
        """
        Adds a single list item (ordered or unordered) to the document with specified indentation.
        Args:
            text: The text content of the list item.
            indent_level: The calculated indentation level (0 for base, 1 for next, etc.).
            is_ordered: True if the item is part of an ordered list.
            item_number: The number for an ordered list item.
        """
        processed_text = self._preprocess_line_content(text)

        prefix = "• " # Default bullet for unordered lists
        if is_ordered:
            num_str = f"{item_number}." if item_number is not None else ""
            prefix = f"{num_str} " if num_str else "• " # Fallback if number missing

        paragraph_text = f"{prefix}{processed_text}"
        p = self.doc.add_paragraph(paragraph_text)

        if indent_level > 0:
            try:
                # Ensure paragraph_format.left_indent is initialized if None
                current_indent = p.paragraph_format.left_indent
                if current_indent is None:
                    current_indent = Inches(0)
                p.paragraph_format.left_indent = current_indent + (LIST_INDENT_PER_LEVEL * indent_level)
            except Exception as e:
                print(f"Warning: Could not apply list indentation for '{processed_text}'. Error: {e}")

    def _consume_unordered_list_block(self, all_lines: List[str], current_index: int) -> int:
        """
        Consumes a block of unordered list items, handling nesting via indentation.
        Args:
            all_lines: The list of all raw lines from the Markdown file.
            current_index: The index of the first line of the list block.
        Returns:
            The index of the line immediately after the processed list block.
        """
        idx = current_index
        # For relative indentation, could track base_indent_spaces of the current list block
        # For now, using absolute indentation based on spaces.

        while idx < len(all_lines):
            raw_line = all_lines[idx].rstrip('\n')
            line_content_stripped = raw_line.strip()

            item_text_after_marker: Optional[str] = None

            for marker_with_space in UNORDERED_LIST_MARKERS:
                if line_content_stripped.startswith(marker_with_space):
                    item_text_after_marker = line_content_stripped[len(marker_with_space):]
                    break

            if item_text_after_marker is None: # Not an unordered list item, so block ends
                break

            leading_spaces = len(raw_line) - len(raw_line.lstrip(' '))
            # Simple heuristic: 2 spaces per indent level.
            # A more robust system might establish a base indent for the list.
            indent_level = leading_spaces // 2 # Integer division

            self._add_list_item(item_text_after_marker, indent_level, is_ordered=False)
            idx += 1

        return idx

    def _consume_ordered_list_block(self, all_lines: List[str], current_index: int) -> int:
        """
        Consumes a block of ordered list items.
        Args:
            all_lines: The list of all raw lines from the Markdown file.
            current_index: The index of the first line of the ordered list.
        Returns:
            The index of the line immediately after the processed list block.
        """
        idx = current_index
        expected_list_number = 0 # Can be used to track numbering continuity

        while idx < len(all_lines):
            raw_line = all_lines[idx].rstrip('\n')
            stripped_line = raw_line.strip()

            match = ORDERED_LIST_PATTERN.match(stripped_line)
            if match:
                item_text = stripped_line[match.end():] # Text after "N. "

                parsed_num_str = stripped_line[:match.start(0) + stripped_line[match.start(0):].find('.')]
                current_item_number = None
                try:
                    parsed_num = int(parsed_num_str)
                    if expected_list_number == 0 or parsed_num == 1: # Start of a new list or explicit "1."
                        expected_list_number = parsed_num
                    elif parsed_num == expected_list_number +1: # Continues sequence
                         expected_list_number = parsed_num
                    else: # Sequence broken, treat as new start or just use its number
                        expected_list_number = parsed_num # Reset expected number
                    current_item_number = expected_list_number
                except ValueError:
                    # If number parsing fails, increment or handle as an issue
                    expected_list_number +=1
                    current_item_number = expected_list_number # Fallback

                leading_spaces = len(raw_line) - len(raw_line.lstrip(' '))
                indent_level = leading_spaces // 2 # Basic indent level for ordered lists

                self._add_list_item(item_text, indent_level, is_ordered=True, item_number=current_item_number)
                idx += 1
            else:
                break # Not a conforming ordered list item
        return idx

    def _add_paragraph_text(self, line_text: str):
        """
        Adds a regular paragraph to the document after preprocessing.
        Args:
            line_text: The raw text of the paragraph.
        """
        processed_line = self._preprocess_line_content(line_text)
        if processed_line:
            self.doc.add_paragraph(processed_line)
        # If line_text was not empty but processed_line is (e.g. just whitespace or empty anchor),
        # an empty paragraph is not added. This matches original behavior.
        # To add empty paragraphs for lines that were originally blank (not just whitespace),
        # that check would need to be in the main loop before stripping.

    def convert_lines_to_doc(self, markdown_content_lines: List[str]):
        """
        Processes a list of markdown lines and populates the self.doc object.
        This is the main dispatch loop for parsing Markdown elements.
        Args:
            markdown_content_lines: A list of strings, where each string is a line from the Markdown source.
        """
        self.doc = Document()  # Ensure a fresh document for each conversion

        current_line_idx = 0
        num_lines = len(markdown_content_lines)

        while current_line_idx < num_lines:
            raw_line_original_spacing = markdown_content_lines[current_line_idx].rstrip('\n')
            stripped_line_content = raw_line_original_spacing.strip()

            if not stripped_line_content:  # Line is empty or contains only whitespace
                # Optionally, add an empty paragraph for visual spacing:
                # self.doc.add_paragraph("")
                current_line_idx += 1
                continue

            # --- Block Element Detection (Order is important) ---

            # 1. Fenced Code Blocks (```)
            if stripped_line_content.startswith(FENCED_CODE_BLOCK_MARKER):
                current_line_idx = self._consume_fenced_code_block(markdown_content_lines, current_line_idx)

            # 2. Markdown Headings (#, ##, etc.)
            elif stripped_line_content.startswith(HEADING_MARKER):
                level = 0
                temp_line = stripped_line_content
                # Count '#' characters for the heading level
                while temp_line.startswith(HEADING_MARKER):
                    level += 1
                    temp_line = temp_line[len(HEADING_MARKER):]

                heading_text = stripped_line_content[level:].strip() # Text after all '#'
                self._add_heading(heading_text, markdown_level=level, explicit_docx_level=2)
                current_line_idx += 1

            # 3. Heuristic Section Titles (Args:, Returns:)
            elif stripped_line_content.lower() in SECTION_TITLE_KEYWORDS:
                self._add_heading(stripped_line_content.capitalize(), explicit_docx_level=ASSUMED_SECTION_HEADING_LEVEL)
                current_line_idx += 1

            # 4. Python Function Definitions (heuristic based on 'def ...):')
            elif PYTHON_DEF_PATTERN.match(stripped_line_content):
                p = self.doc.add_paragraph()
                # Pass raw line to preserve original indentation for code display
                self._format_code_paragraph(p, raw_line_original_spacing, is_block=False)
                current_line_idx += 1

            # 5. Unordered Lists (*, -, •)
            elif any(stripped_line_content.startswith(marker) for marker in UNORDERED_LIST_MARKERS):
                 current_line_idx = self._consume_unordered_list_block(markdown_content_lines, current_line_idx)

            # 6. Ordered Lists (1., 2., etc.)
            elif ORDERED_LIST_PATTERN.match(stripped_line_content):
                current_line_idx = self._consume_ordered_list_block(markdown_content_lines, current_line_idx)

            # 7. Default: Treat as a Paragraph
            else:
                # Pass raw line; _add_paragraph_text handles preprocessing
                self._add_paragraph_text(raw_line_original_spacing)
                current_line_idx += 1

    def save_document(self, docx_filepath: Path):
        """
        Saves the generated DOCX document to the specified file path.
        Args:
            docx_filepath: The Path object representing the output DOCX file.
        Raises:
            IOError: If there is an error writing the file.
        """
        try:
            self.doc.save(docx_filepath)
        except IOError as e:
            print(f"Error: Could not write DOCX file to '{docx_filepath}'. Details: {e}")
            raise


def process_single_markdown_file(markdown_filepath: Path, docx_filepath: Path):
    """
    Reads a single Markdown file, converts its content, and saves it as a DOCX document.
    Args:
        markdown_filepath: Path to the input Markdown file.
        docx_filepath: Path to the output DOCX file.
    """
    # print(f"Processing {markdown_filepath.name}...")
    try:
        with open(markdown_filepath, "r", encoding="utf-8") as f:
            markdown_lines = f.readlines()
    except IOError as e:
        print(f"Error: Could not read Markdown file '{markdown_filepath}'. Details: {e}")
        return

    converter = MarkdownToDocxConverter()
    converter.convert_lines_to_doc(markdown_lines)

    try:
        converter.save_document(docx_filepath)
        print(f"Successfully converted: {markdown_filepath.name} -> {docx_filepath.name}")
    except IOError:
        # Error message already printed by save_document
        pass


def convert_md_directory_to_docx(input_dir_str: str, output_dir_str: str):
    """
    Converts all Markdown (.md) files in an input directory to DOCX files
    in an output directory. Creates the output directory if it doesn't exist.
    Args:
        input_dir_str: Path string for the input directory containing .md files.
        output_dir_str: Path string for the output directory where .docx files will be saved.
    """
    input_dir = Path(input_dir_str)
    output_dir = Path(output_dir_str)

    if not input_dir.is_dir():
        print(f"Error: Input directory '{input_dir}' does not exist or is not a directory.")
        return

    try:
        output_dir.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        print(f"Error: Could not create output directory '{output_dir}'. Details: {e}")
        return

    markdown_files_found = False
    for md_file_path in input_dir.glob("*.md"):
        markdown_files_found = True
        docx_filename = md_file_path.stem + ".docx"
        docx_filepath = output_dir / docx_filename

        process_single_markdown_file(md_file_path, docx_filepath)

    if not markdown_files_found:
        print(f"No Markdown (.md) files found in '{input_dir}'.")



def authenticate():
    """Authenticate using service account."""
    creds_dict = json.loads(os.environ["GOOGLE_APPLICATION_CREDENTIALS"])
    creds = service_account.Credentials.from_service_account_info(
        creds_dict,
        scopes=["https://www.googleapis.com/auth/drive"]
    )
    return build("drive", "v3", credentials=creds)

def get_folder_id_by_name(service, folder_name: str, is_shared_drive: bool = False) -> str | None:
    """
    Searches for a folder by its name in Google Drive (or Shared Drives) and returns its ID.

    Args:
        service: The Google Drive API service object.
        folder_name (str): The name of the folder to search for.
        is_shared_drive (bool): Set to True if searching within a Shared Drive.

    Returns:
        str | None: The ID of the folder if found, otherwise None.
    """
    try:
        query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"

        # Add supportsAllDrives=True for searching in Shared Drives
        response = service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name)',
            corpora='allDrives' if is_shared_drive else 'user', # Search all drives for Shared Drives, else user's drive
            includeItemsFromAllDrives=True if is_shared_drive else False, # Required for allDrives
            supportsAllDrives=True if is_shared_drive else False # Required for allDrives
        ).execute()

        folders = response.get('files', [])
        if folders:
            print(f"Found folder '{folder_name}' with ID: {folders[0]['id']}")
            return folders[0]['id']
        else:
            print(f"Folder '{folder_name}' not found in Google Drive.")
            return None
    except Exception as e:
        print(f"Error searching for folder '{folder_name}': {e}")
        return None

def upload_docx_to_gdoc(service, file_path: str, folder_id: str = None) -> str | None:
    """
    Uploads a .docx file to Google Drive and converts it to a Google Doc.
    Checks if a file with the same name already exists in the target folder.
    If it exists, the existing file is updated; otherwise, a new file is created.

    Args:
        service: The Google Drive API service object.
        file_path (str): The path to the .docx file to upload.
        folder_id (str, optional): The ID of the target folder in Google Drive.
                                   If None, the file will be uploaded to the root folder.

    Returns:
        str | None: The ID of the uploaded or updated Google Doc if successful, otherwise None.
    """
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None

    if not file_path.lower().endswith('.docx'):
        print(f"Warning: File '{file_path}' is not a .docx file. Attempting upload but conversion might fail.")

    gdoc_title = os.path.splitext(os.path.basename(file_path))[0]

    file_metadata = {
        'name': gdoc_title,
        'mimeType': 'application/vnd.google-apps.document' # This is the target MIME type (Google Doc)
    }
    # If a folder_id is provided, the new file (if created) will be placed there.
    # For searching existing files, the query also considers the folder.

    media_body = MediaFileUpload(
        file_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', # Source MIME type
        resumable=True
    )

    existing_file_id = None
    try:
        # 1. Search for an existing Google Doc with the same name in the target folder
        query_parts = [
            f"name='{gdoc_title}'",
            f"mimeType='{file_metadata['mimeType']}'", # Search for Google Docs
            f"'{folder_id}' in parents", # Search only in target folder
            "trashed=false"
        ]
        search_query = " and ".join(query_parts)

        print(f"Searching for existing Google Doc '{gdoc_title}' in Drive folder ID: {folder_id}...")
        response = service.files().list(
            q=search_query,
            spaces='drive',
            fields='files(id, name)',
            # Add these parameters to support Shared Drives
            corpora='allDrives',
            includeItemsFromAllDrives=True,
            supportsAllDrives=True
        ).execute()

        files = response.get('files', [])
        if files:
            existing_file_id = files[0]['id']
            print(f"Found existing Google Doc '{gdoc_title}' (ID: {existing_file_id}). Attempting to update it.")
    except Exception as e:
        print(f"Error searching for existing file '{gdoc_title}': {e}")
        # If search fails, proceed as if no file was found (attempt to create a new one)

    try:
        if existing_file_id:
            # Update existing file
            updated_file = service.files().update(
                fileId=existing_file_id,
                media_body=media_body, # New content from local .docx
                fields='id, name',
                # Add these parameters to support Shared Drives
                supportsAllDrives=True
            ).execute()
            print(f"Successfully updated '{updated_file.get('name')}' (ID: {updated_file.get('id')}).")
            return updated_file.get('id')
        else:
            # Create new file
            print(f"No existing file '{gdoc_title}' found in target folder. Creating a new one.")
            # For creating, we explicitly add the parent folder
            file_metadata['parents'] = [folder_id]
            uploaded_file = service.files().create(
                body=file_metadata, # Contains name and target mimeType
                media_body=media_body, # Contains actual file content and source mimeType
                fields='id, name',
                # Add these parameters to support Shared Drives
                supportsAllDrives=True
            ).execute()
            print(f"Successfully uploaded '{uploaded_file.get('name')}' (ID: {uploaded_file.get('id')}) to Google Drive.")
            return uploaded_file.get('id')
    except Exception as e:
        print(f"Failed to upload/update '{file_path}': {e}")
        return None

def main(local_source_folder: str, drive_target_folder_id: str):
    """
    Main function to handle authentication, get service, find and upload
    all DOCX files from a specified local folder to a specified Google Drive folder.

    Args:
        local_source_folder (str): The path to the local folder containing DOCX files.
        drive_target_folder_id (str): The ID of the Google Drive folder to upload to.
    """
    service = authenticate()
    \
    # Validate local source folder
    if not os.path.isdir(local_source_folder):
        print(f"Local source folder not found or is not a directory: {local_source_folder}. Exiting.")
        return

    print(f"Scanning for .docx files in local folder: {local_source_folder}")
    docx_files_found = []
    for filename in os.listdir(local_source_folder):
        # Check if it's a file and ends with .docx (case-insensitive)
        if os.path.isfile(os.path.join(local_source_folder, filename)) and filename.lower().endswith('.docx'):
            full_path = os.path.join(local_source_folder, filename)
            docx_files_found.append(full_path)

    if not docx_files_found:
        print(f"No .docx files found in '{local_source_folder}'.")
        return

    print(f"Found {len(docx_files_found)} .docx files to upload.")
    print(f"Target Google Drive Folder ID: {drive_target_folder_id}")

    uploaded_ids = []
    for docx_file_path in docx_files_found:
        uploaded_id = upload_docx_to_gdoc(service, docx_file_path, drive_target_folder_id)
        if uploaded_id:
            uploaded_ids.append((docx_file_path, uploaded_id))

    if uploaded_ids:
        print("\n--- Upload Summary ---")
        for original_path, uploaded_id in uploaded_ids:
            print(f"'{original_path}' uploaded/updated as Google Doc with ID: {uploaded_id}")
    else:
        print("\nNo files were successfully uploaded or updated.")



#%%
if __name__ == "__main__":
    MD_FOLDER = './Mds'
    DOCX_FOLDER = './Docx'

    input_dir = MD_FOLDER
    output_dir = DOCX_FOLDER
    convert_md_directory_to_docx(input_dir, output_dir)

    LOCAL_DOCS_FOLDER_PATH = DOCX_FOLDER

    # REPLACE 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE' with the actual ID of your Google Drive target folder.
    DRIVE_UPLOAD_FOLDER_ID = '1Zp5DcqflMOTzPwpl3ql54GDpIoIy-OSd'

    # Call the main function with the specified local folder and Drive folder ID
    if DRIVE_UPLOAD_FOLDER_ID == 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE':
        print("Critical: Please replace 'YOUR_GOOGLE_DRIVE_FOLDER_ID_HERE' with your actual Google Drive folder ID.")
    else:
        main(LOCAL_DOCS_FOLDER_PATH, DRIVE_UPLOAD_FOLDER_ID)

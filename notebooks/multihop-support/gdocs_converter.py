import os
import base64
import hashlib
from datetime import datetime, UTC
from docx import Document


def get_content_from_local_file(local_path, file_extension):
    """Reads content from a local file based on its extension."""
    # Normalize to lowercase for case-insensitive matching
    ext_lower = file_extension.lower()
    
    text_extensions = ['.txt', '.json', '.log', '.html', '.css', '.py', '.md', '.notebook', '.htm', '.csv', '.ini']
    base64_extensions = ['.webp', '.xlt']
    # Special handling for ambiguous extensions
    ambiguous_extensions = ['.download']  # Could be text or binary
    # System files to ignore content
    system_extensions = ['.lnk', '.ds_store']
    doc_extensions = ['.docx', '.doc']
    
    try:
        if ext_lower in text_extensions:
            with open(local_path, 'r', encoding='utf-8', errors='ignore') as f:
                content_data = f.read()
            
            # Convert text content to Google Docs format (list of paragraph elements)
            paragraphs = content_data.split('\n')
            content = []
            for i, para in enumerate(paragraphs):
                if para.strip():  # Skip empty paragraphs
                    content.append({
                        "elementId": f"p{i+1}",
                        "text": para
                    })
            
        elif ext_lower in doc_extensions:
            # Try to parse as docx using python-docx
            try:
                doc = Document(local_path)
                content = []
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:
                        content.append({
                            "elementId": f"p{i+1}",
                            "text": text
                        })
                
                # Extract tables as text representation
                for t_idx, table in enumerate(doc.tables):
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text:
                            table_text.append(row_text)
                    
                    if table_text:
                        content.append({
                            "elementId": f"t{t_idx+1}",
                            "text": "\n".join(table_text)
                        })
            except Exception as e:
                # Fallback: try to read as plain text
                try:
                    with open(local_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content_data = f.read()
                    
                    paragraphs = content_data.split('\n')
                    content = []
                    for i, para in enumerate(paragraphs):
                        if para.strip():  # Skip empty paragraphs
                            content.append({
                                "elementId": f"p{i+1}",
                                "text": para
                            })
                except Exception as e2:
                    content = []  # Return empty content on error
            
        elif ext_lower in base64_extensions or ext_lower in system_extensions or ext_lower in ambiguous_extensions:
            # For binary or system files, create a single paragraph with file info
            content = [{
                "elementId": "p1",
                "text": f"Binary or system file: {os.path.basename(local_path)}"
            }]
        else:
            # Unknown extension - try to detect if it's text
            try:
                with open(local_path, 'r', encoding='utf-8', errors='strict') as f:
                    content_data = f.read()
                
                paragraphs = content_data.split('\n')
                content = []
                for i, para in enumerate(paragraphs):
                    if para.strip():  # Skip empty paragraphs
                        content.append({
                            "elementId": f"p{i+1}",
                            "text": para
                        })
            except UnicodeDecodeError:
                # Binary file - create a placeholder paragraph
                content = [{
                    "elementId": "p1",
                    "text": f"Binary file: {os.path.basename(local_path)}"
                }]
            
    except Exception as e:
        content = [{
            "elementId": "p1",
            "text": f"Error reading content from {os.path.basename(local_path)}: {e}"
        }]
    
    return content


def convert_doc_to_gdoc_format(file_path, file_name=None):
    """Convert a local file to Google Drive format JSON."""
    if file_name is None:
        file_name = os.path.basename(file_path)
    
    file_extension = os.path.splitext(file_name)[1].lower()
    file_stats = os.stat(file_path)
    file_id = f"file_{hashlib.md5(file_path.encode()).hexdigest()}"
    current_time = datetime.fromtimestamp(file_stats.st_mtime, UTC).strftime('%Y-%m-%dT%H:%M:%SZ')
    
    # Create base metadata matching GDriveDefaultDB.json schema exactly
    json_data = {
        "id": file_id,
        "driveId": "",
        "name": file_name,
        "mimeType": _get_mime_type(file_extension),
        "createdTime": datetime.fromtimestamp(file_stats.st_ctime, UTC).strftime('%Y-%m-%dT%H:%M:%SZ'),
        "modifiedTime": current_time,
        "trashed": False,
        "starred": False,
        "parents": [],
        "owners": ["john.doe@gmail.com"],
        "size": str(file_stats.st_size),
        "permissions": [
            {
                "id": f"permission_{file_id}",
                "role": "owner",
                "type": "user",
                "emailAddress": "john.doe@gmail.com"
            }
        ],
        "suggestionsViewMode": "DEFAULT",
        "includeTabsContent": False,
        "tabs": []
    }
    
    # Get file content in Google Docs format
    content = get_content_from_local_file(file_path, file_extension)
    json_data["content"] = content
    
    # Add revisions if we have content
    if content:
        # Create a simplified revision entry
        json_data["revisions"] = [
            {
                "id": "rev-1",
                "mimeType": json_data["mimeType"],
                "modifiedTime": current_time,
                "keepForever": False,
                "originalFilename": file_name,
                "size": str(file_stats.st_size)
            }
        ]
        
        # Add export formats for supported types
        if json_data["mimeType"] in ["application/vnd.google-apps.document", "text/plain"]:
            json_data["exportFormats"] = {
                "application/pdf": content[:50] + "...",  # Truncated for demo
                "application/msword": content[:50] + "..."
            }
        elif json_data["mimeType"] in ["application/vnd.google-apps.spreadsheet", "application/vnd.ms-excel"]:
            json_data["exportFormats"] = {
                "application/pdf": content[:50] + "...",
                "application/vnd.ms-excel": content[:50] + "..."
            }
    else:
        # For files without content, ensure empty structures exist
        json_data["revisions"] = []
        json_data["exportFormats"] = {}
    
    return json_data


def _get_mime_type(file_extension):
    """Get MIME type based on file extension."""
    # Normalize to lowercase for case-insensitive matching
    ext_lower = file_extension.lower()
    
    mime_types = {
        '.txt': 'text/plain',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.css': 'text/css',
        '.js': 'application/javascript',
        '.json': 'application/json',
        '.csv': 'text/csv',
        '.md': 'text/markdown',
        '.py': 'text/x-python',
        '.pdf': 'application/pdf',
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.webp': 'image/webp',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xls': 'application/vnd.ms-excel',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.ppt': 'application/vnd.ms-powerpoint',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        '.xlt': 'application/vnd.ms-excel',
        '.ini': 'text/plain',
        '.log': 'text/plain',
        '.notebook': 'application/json',
        '.download': 'application/octet-stream',
        '.lnk': 'application/x-ms-shortcut',
        '.ds_store': 'application/octet-stream'
    }
    return mime_types.get(ext_lower, 'application/octet-stream') 